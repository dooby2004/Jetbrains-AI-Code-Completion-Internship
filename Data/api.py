from datetime import date, datetime
import hashlib
from io import BytesIO
import os
import shutil
from pathlib import Path
from uuid import uuid4

import docx
import numpy as np
import torch
from fastapi import APIRouter, HTTPException, UploadFile
from PyPDF2 import PdfReader
from transformers import AutoConfig, AutoModelForSequenceClassification, AutoTokenizer

from server.config import Config
from server.db import DBService
from server.models import Document

router = APIRouter()

# UPLOAD_FOLDER = "uploads"
# Path(
#     SERVER_ROOT_PATH / UPLOAD_FOLDER,
# ).mkdir(parents=True, exist_ok=True)

# Validates Upload Types
VALID_UPLOAD_TYPES = [
    "text/plain",
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.wordproccessingml.document",
]

ig = AutoConfig.from_pretrained(Config.SERVER_ROOT_PATH / "server/model")
tokenizer = AutoTokenizer.from_pretrained(Config.SERVER_ROOT_PATH / "server/model")
model = AutoModelForSequenceClassification.from_pretrained(
    Config.SERVER_ROOT_PATH / "server/model"
)


# DBService.instance().db.connect()
# DBService.instance().db.create_tables([Person])
# uncle_bob = Person(name="Bob", birthday=date(1960, 1, 15))
# uncle_bob.save()


@router.get("/documents")
def retrieve_docs():
    documents = []

    # for filename in os.listdir(Config.UPLOAD_PATH):
    #     file_path = os.path.join(Config.UPLOAD_PATH, filename)
    #     if os.path.isfile(file_path):
    #         parts = filename.split(".")
    #         hash_part = parts[0]
    #         extension_part = parts[-1]
    #         name_part = ".".join(parts[1:-1])

    #         document = {
    #             "hash": hash_part,
    #             "name": name_part,
    #             "extension": extension_part,
    #         }
    #         documents.append(document)

    for document in Document.select():
        documents.append({
            "name": document.name,
            "hash": document.hash,
            "content": document.content,
            "prediction": document.prediction,
            "uploaded_at": document.uploaded_at,
            "last_predicted_at": document.last_predicted_at,
        })

    return documents


def get_file_hash(file_bin):
    file_hash = hashlib.sha256()

    # # for chunk in iter(lambda: file.file.read(4096), b""):
    file_hash.update(file_bin)

    return file_hash.hexdigest()


@router.post("/upload")
def upload_file(file: UploadFile):
    if file.content_type not in VALID_UPLOAD_TYPES:
        raise HTTPException(status_code=400, detail="invalid content type")

    file_bin = file.file.read()

    file_hash = get_file_hash(file_bin)

    # file_extension = file.filename.split(".")[-1]

    # text = file.file.read().decode("utf-8")

    text = ""
    if file.content_type == "text/plain":
        text = file_bin.decode("utf-8")

    elif file.content_type == "application/pdf":
        file.file.seek(0)  # Reset the file pointer
        pdf = PdfReader(BytesIO(file_bin))
        text = "\n".join(page.extract_text() for page in pdf.pages)

    elif file.content_type in [
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/vnd.openxmlformats-officedocument.wordproccessingml.document",
    ]:
        in_memory_file = BytesIO(file_bin)
        doc = docx.Document(in_memory_file)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)


    current_time = datetime.now().time()

    # file_path = Config.UPLOAD_PATH / f"{file_hash}.{file.filename}.{file_extension}"

    if DBService.instance().db.is_closed():
        DBService.instance().db.connect()
    
    DBService.instance().db.create_tables([Document], safe = True)
    document = Document(name = file.filename, hash = file_hash, content = text, uploaded_at = current_time)
    document.save()

    DBService.instance().db.close()

    # text = ""

    # if file.content_type == "text/plain":
    #     text = "".join([line.decode("utf-8") for line in file.file.readlines()])

    # elif file.content_type == "application/pdf":
    #     reader = PdfReader(file.file)
    #     for page in reader.pages:
    #         text += page.extract_text()

    # elif file.content_type in [
    #     "application/msword",
    #     "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    #     "application/vnd.openxmlformats-officedocument.wordproccessingml.document",
    # ]:
    # name = str(uuid4())
    # with open(Config.SERVER_ROOT_PATH / f"/{name}", "wb") as f:
    #     shutil.copyfileobj(file.file, f)
    # doc = docx.Document(Config.SERVER_ROOT_PATH / f"/{name}")
    # Path(f"./{name}").unlink()
    # for paragraph in doc.paragraphs:
    #     text += paragraph.text

    # if os.path.exists(file_path):
    #     os.remove(file_path)

    # with open(file_path, "wb") as f:
    #     f.write(file_bin)
        # shutil.copyfileobj(file_bin, f)

    # doc = docx.Document(file_path)

    # for paragraph in doc.paragraphs:
    #     text += paragraph.text

    # else:
    #     text = file.file.read().decode("utf-8")


    # encoding_input = tokenizer(text, return_tensors="pt")
    # outputs = model(**encoding_input)
    # logits = outputs.logits

    # sigmoid = torch.nn.Sigmoid()
    # probs = sigmoid(logits.squeeze().cpu())
    # predictions = np.zeros(probs.shape)
    # predictions[np.where(probs >= 0.5)] = 1

    # predicted_labels = [
    #     config.id2label[idx] for idx, label in enumerate(predictions) if label == 1.0
    # ]

    return {
        # "content": text,
        # "predictions": predicted_labels,
        # "type": file.content_type,
        "file_hash": file_hash,
        "file_name": file.filename,
    }
